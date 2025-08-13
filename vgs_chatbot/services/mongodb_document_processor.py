"""MongoDB-based document processor service implementation with RAG pipeline."""

import uuid
from datetime import UTC, datetime
from typing import Any

import openpyxl
import tiktoken
from docx import Document as DocxDocument
from pypdf import PdfReader
from sentence_transformers import SentenceTransformer

from vgs_chatbot.interfaces.document_processor_interface import (
    DocumentProcessorInterface,
)
from vgs_chatbot.models.document import Document, ProcessedDocument
from vgs_chatbot.models.vector_document import DocumentSummary, VectorDocument
from vgs_chatbot.repositories.vector_repository import VectorRepository
from vgs_chatbot.utils.manifest import DocumentManifest
from vgs_chatbot.utils.retrieval_enhancer import RetrievalEnhancer


class MongoDBDocumentProcessor(DocumentProcessorInterface):
    """MongoDB-based RAG document processor implementation."""

    def __init__(
        self,
        vector_repository: VectorRepository,
        embedding_model: str = "BAAI/bge-small-en-v1.5",
        manifest_path: str = "data/vectors/manifest.json",
    ) -> None:
        """Initialize document processor.

        Args:
            vector_repository: MongoDB vector repository
            embedding_model: Name of sentence transformer model for embeddings
            manifest_path: Path to manifest file for tracking processed documents
        """
        self.vector_repository = vector_repository
        self.embedding_model_name = embedding_model

        # Try preferred model with fallback
        try:
            self.embedding_model = SentenceTransformer(embedding_model)
            print(f"Loaded embedding model: {embedding_model}")
        except Exception as e:
            print(
                f"Failed to load {embedding_model}, falling back to multi-qa-MiniLM-L6-cos-v1: {e}"
            )
            self.embedding_model_name = "multi-qa-MiniLM-L6-cos-v1"
            self.embedding_model = SentenceTransformer("multi-qa-MiniLM-L6-cos-v1")

        self.manifest = DocumentManifest(manifest_path)
        self.retrieval_enhancer = RetrievalEnhancer()

        # Section-based chunking parameters
        self.chunk_target_tokens = 450  # Target tokens per chunk
        self.chunk_max_tokens = 1000  # Maximum tokens for large sections

        # Initialize tiktoken encoder for accurate token counting
        self.tokenizer = tiktoken.get_encoding(
            "cl100k_base"
        )  # GPT-4 compatible encoding

        # Determine model-specific formatting needs
        self._requires_prefixes = self._model_requires_prefixes(
            self.embedding_model_name
        )
        if self._requires_prefixes:
            print(f"Model {self.embedding_model_name} requires query/passage prefixes")

        # Deduplication tracking
        self._chunk_hashes: set[str] = set()

        # BM25 index for hybrid retrieval (initialized when documents are indexed)
        self._bm25_index = None
        self._bm25_corpus = []
        self._bm25_ids = []

        # Initialize vector search index
        self._setup_vector_search_index()

    def _setup_vector_search_index(self) -> None:
        """Set up MongoDB Atlas Vector Search index."""
        try:
            # Determine embedding dimensions based on the model
            if "bge-small" in self.embedding_model_name:
                dimensions = 384
            elif "multi-qa-MiniLM-L6" in self.embedding_model_name:
                dimensions = 384
            else:
                # Default dimensions, may need adjustment based on your model
                dimensions = 384

            # Try to create or verify vector search index
            self.vector_repository.create_vector_search_index(
                index_name="vector_index",
                vector_field="embedding",
                similarity_function="cosine",
                dimensions=dimensions
            )
        except Exception as e:
            print(f"Warning: Could not set up vector search index: {e}")
            print("The system will use fallback similarity search.")

    def _token_count(self, text: str) -> int:
        """Count tokens in text using tiktoken encoder.

        Args:
            text: Text to count tokens for

        Returns:
            Number of tokens
        """
        return len(self.tokenizer.encode(text))

    def _model_requires_prefixes(self, model_name: str) -> bool:
        """Check if model requires query/passage prefixes.

        Args:
            model_name: Name of the embedding model

        Returns:
            True if prefixes are required
        """
        prefix_models = [
            "sentence-transformers/multi-qa-MiniLM-L6-cos-v1",
            "multi-qa-MiniLM-L6-cos-v1",
        ]
        return any(prefix in model_name for prefix in prefix_models)

    def _format_text_for_embedding(self, text: str, is_query: bool = False) -> str:
        """Format text for embedding with model-specific prefixes if needed.

        Args:
            text: Text to format
            is_query: Whether this is a query (vs passage)

        Returns:
            Formatted text
        """
        if not self._requires_prefixes:
            return text

        if is_query:
            return f"query: {text}"
        return f"passage: {text}"

    def add_document(self, document: Document) -> ProcessedDocument:
        """Add document to knowledge base with embeddings.

        Args:
            document: Document to process and add

        Returns:
            Processed document with metadata

        Raises:
            Exception: If document processing fails
        """
        try:
            # Extract text content based on file type
            content = self._extract_text(document)

            if not content.strip():
                raise ValueError(f"No extractable content from {document.name}")

            # Generate document ID
            doc_id = document.id or str(uuid.uuid4())

            # Clear any existing chunks for this document
            self.vector_repository.delete_document_chunks(doc_id)

            # Create processed document
            processed_doc = ProcessedDocument(
                id=doc_id,
                original_document=document,
                content=content,
                chunks=[],  # Will be populated by chunking
                processed_at=datetime.now(UTC),
                metadata={
                    "file_type": document.file_type,
                    "size": document.size,
                    "modified_date": document.modified_date,
                }
            )

            # Perform intelligent chunking
            chunks = self._intelligent_chunk_document(content, doc_id)
            processed_doc.chunks = [chunk.content for chunk in chunks]

            # Store chunks in MongoDB
            self._store_document_chunks(chunks, processed_doc)

            # Generate and store document summary if content is substantial
            if len(content) > 2000:
                self._store_document_summary(processed_doc)

            # Update manifest
            self.manifest.update_document_entry(
                document, len(chunks), self.embedding_model_name
            )

            print(f"✅ Added document '{document.name}' with {len(chunks)} chunks")
            return processed_doc

        except Exception as e:
            print(f"❌ Error processing document {document.name}: {e}")
            raise

    def _store_document_chunks(self, chunks: list[VectorDocument], processed_doc: ProcessedDocument) -> None:
        """Store document chunks with embeddings in MongoDB.

        Args:
            chunks: List of document chunks to store
            processed_doc: The processed document
        """
        for chunk in chunks:
            # Generate embedding for chunk content
            formatted_content = self._format_text_for_embedding(chunk.content)
            embedding = self.embedding_model.encode(formatted_content).tolist()

            # Update chunk with embedding and add document name to metadata
            chunk.embedding = embedding
            chunk.metadata["document_name"] = processed_doc.original_document.name
            chunk.metadata["file_type"] = processed_doc.original_document.file_type

            # Store in MongoDB
            self.vector_repository.add_document_chunk(chunk)

    def _store_document_summary(self, processed_doc: ProcessedDocument) -> None:
        """Store document summary with embedding in MongoDB.

        Args:
            processed_doc: The processed document
        """
        # Generate summary (first 500 words as a simple summary)
        words = processed_doc.content.split()
        summary_text = " ".join(words[:500])

        # Generate embedding for summary
        formatted_summary = self._format_text_for_embedding(summary_text)
        summary_embedding = self.embedding_model.encode(formatted_summary).tolist()

        # Create and store summary document
        doc_summary = DocumentSummary(
            document_id=processed_doc.id,
            summary_text=summary_text,
            embedding=summary_embedding,
            metadata={
                "document_name": processed_doc.original_document.name,
                "is_summary": True,
                **processed_doc.metadata
            }
        )

        self.vector_repository.add_document_summary(doc_summary)

    def _intelligent_chunk_document(self, content: str, doc_id: str) -> list[VectorDocument]:
        """Intelligently chunk document content preserving section boundaries.

        Args:
            content: Full document content
            doc_id: Document ID

        Returns:
            List of vector document chunks
        """
        chunks = []

        # Identify document sections
        sections = self._identify_document_sections(content)

        chunk_index = 0
        for section in sections:
            section_title = self._extract_section_title(section)
            section_chunks = self._chunk_section(section, doc_id, chunk_index, section_title)
            chunks.extend(section_chunks)
            chunk_index += len(section_chunks)

        return chunks

    def _identify_document_sections(self, text: str) -> list[str]:
        """Identify document sections based on structural patterns.

        Args:
            text: Document text to analyze

        Returns:
            List of identified sections
        """
        lines = text.split('\n')
        sections = []
        current_section = []

        for i, line in enumerate(lines):
            line = line.strip()

            # Check if this line starts a new section
            if self._is_section_boundary(line, i, lines) and current_section:
                # Save current section
                sections.append('\n'.join(current_section))
                current_section = [line]
            else:
                current_section.append(line)

        # Add final section
        if current_section:
            sections.append('\n'.join(current_section))

        # Fallback: if no clear sections found, split by pages or size
        if len(sections) <= 1:
            sections = self._split_by_page_markers(text)

        return sections

    def _is_section_boundary(self, line: str, line_idx: int, all_lines: list[str]) -> bool:
        """Determine if a line represents a section boundary.

        Args:
            line: Current line to check
            line_idx: Index of current line
            all_lines: All lines in the document

        Returns:
            True if this line starts a new section
        """
        if not line:
            return False

        # Common section patterns
        section_patterns = [
            # Numbered sections: "1.", "1.1", "Section 1"
            r'^\d+\.(\d+\.)*\s',
            # Lettered sections: "A.", "a)"
            r'^[A-Z]\.\s|^[a-z]\)\s',
            # Word sections: "SECTION", "CHAPTER", "PART"
            r'^(SECTION|CHAPTER|PART|APPENDIX)\s+[A-Z0-9]',
            # Uppercase headers
            r'^[A-Z][A-Z\s]{10,}$',
            # Line with only numbers and periods
            r'^\d+(\.\d+)*\s*$'
        ]

        import re
        return any(re.match(pattern, line) for pattern in section_patterns)

    def _extract_section_title(self, section: str) -> str | None:
        """Extract title from section content.

        Args:
            section: Section content

        Returns:
            Section title or None
        """
        lines = section.strip().split('\n')
        if lines:
            first_line = lines[0].strip()
            if len(first_line) < 100 and first_line:  # Reasonable title length
                return first_line
        return None

    def _split_by_page_markers(self, text: str) -> list[str]:
        """Split text by page markers as fallback method.

        Args:
            text: Text to split

        Returns:
            List of page-based sections
        """
        # Look for common page markers
        import re
        page_markers = [
            r'\n\s*Page\s+\d+\s*\n',
            r'\n\s*-\s*\d+\s*-\s*\n',
            r'\n\s*\|\s*\d+\s*\|\s*\n'
        ]

        for pattern in page_markers:
            sections = re.split(pattern, text)
            if len(sections) > 1:
                return [s.strip() for s in sections if s.strip()]

        # Ultimate fallback: split by token count
        return self._split_by_token_count(text)

    def _split_by_token_count(self, text: str) -> list[str]:
        """Split text by token count as ultimate fallback.

        Args:
            text: Text to split

        Returns:
            List of token-based sections
        """
        words = text.split()
        sections = []
        current_section = []
        current_tokens = 0

        for word in words:
            word_tokens = self._token_count(word)
            if current_tokens + word_tokens > self.chunk_target_tokens and current_section:
                sections.append(' '.join(current_section))
                current_section = [word]
                current_tokens = word_tokens
            else:
                current_section.append(word)
                current_tokens += word_tokens

        if current_section:
            sections.append(' '.join(current_section))

        return sections

    def _chunk_section(
        self,
        section: str,
        doc_id: str,
        start_index: int,
        section_title: str | None = None
    ) -> list[VectorDocument]:
        """Chunk a section into appropriately sized pieces.

        Args:
            section: Section content to chunk
            doc_id: Document ID
            start_index: Starting chunk index
            section_title: Title of the section

        Returns:
            List of vector document chunks
        """
        token_count = self._token_count(section)

        if token_count <= self.chunk_max_tokens:
            # Section fits in single chunk
            chunk_id = f"{doc_id}_chunk_{start_index}"
            return [VectorDocument(
                document_id=doc_id,
                chunk_id=chunk_id,
                content=section,
                embedding=[],  # Will be populated later
                chunk_index=start_index,
                section_title=section_title,
                metadata={
                    "token_count": token_count,
                    "section_title": section_title
                }
            )]

        # Split oversized section
        subsections = self._split_oversized_section(section)
        chunks = []

        for i, subsection in enumerate(subsections):
            chunk_id = f"{doc_id}_chunk_{start_index + i}"
            chunks.append(VectorDocument(
                document_id=doc_id,
                chunk_id=chunk_id,
                content=subsection,
                embedding=[],  # Will be populated later
                chunk_index=start_index + i,
                section_title=section_title,
                metadata={
                    "token_count": self._token_count(subsection),
                    "section_title": section_title,
                    "subsection_index": i
                }
            ))

        return chunks

    def _split_oversized_section(self, section: str) -> list[str]:
        """Split an oversized section while preserving subsection boundaries.

        Args:
            section: Section text to split

        Returns:
            List of sub-chunks
        """
        # Try to split by paragraphs first
        paragraphs = section.split('\n\n')

        if len(paragraphs) > 1:
            chunks = []
            current_chunk = []
            current_tokens = 0

            for para in paragraphs:
                para_tokens = self._token_count(para)

                if current_tokens + para_tokens > self.chunk_target_tokens and current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                    current_chunk = [para]
                    current_tokens = para_tokens
                else:
                    current_chunk.append(para)
                    current_tokens += para_tokens

            if current_chunk:
                chunks.append('\n\n'.join(current_chunk))

            return chunks

        # Fallback: split by sentences or words
        return self._split_by_token_count(section)

    def _extract_text(self, document: Document) -> str:
        """Extract text content from document based on file type.

        Args:
            document: Document to extract text from

        Returns:
            Extracted text content

        Raises:
            ValueError: If file type is not supported or content cannot be extracted
        """
        try:
            # Check if document has file content (MongoDB storage) or file path (legacy)
            if document.file_content:
                # Extract from bytes content stored in MongoDB
                if document.file_type == "application/pdf":
                    return self._extract_pdf_text_from_bytes(document.file_content)
                elif document.file_type in [
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ]:
                    return self._extract_docx_text_from_bytes(document.file_content)
                elif document.file_type in [
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                ]:
                    return self._extract_xlsx_text_from_bytes(document.file_content)
                elif document.file_type == "text/plain":
                    return document.file_content.decode('utf-8')
                else:
                    raise ValueError(f"Unsupported file type: {document.file_type}")
            elif document.file_path:
                # Legacy file path extraction
                if document.file_type == "application/pdf":
                    return self._extract_pdf_text(document.file_path)
                elif document.file_type in [
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ]:
                    return self._extract_docx_text(document.file_path)
                elif document.file_type in [
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                ]:
                    return self._extract_xlsx_text(document.file_path)
                elif document.file_type == "text/plain":
                    with open(document.file_path, encoding='utf-8') as f:
                        return f.read()
                else:
                    raise ValueError(f"Unsupported file type: {document.file_type}")
            else:
                raise ValueError(f"Document {document.name} has no file content or file path")
        except Exception as e:
            raise ValueError(f"Failed to extract text from {document.name}: {e}")

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content
        """
        text_content = []
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(f"--- Page {page_num} ---\n{page_text}")
        return "\n\n".join(text_content)

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        doc = DocxDocument(file_path)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    def _extract_xlsx_text(self, file_path: str) -> str:
        """Extract text from XLSX file.

        Args:
            file_path: Path to XLSX file

        Returns:
            Extracted text content
        """
        workbook = openpyxl.load_workbook(file_path)
        text_content = []

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_text = [f"--- Sheet: {sheet_name} ---"]

            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    sheet_text.append(row_text)

            text_content.extend(sheet_text)

        return "\n".join(text_content)

    def _extract_pdf_text_from_bytes(self, file_content: bytes) -> str:
        """Extract text from PDF file content.

        Args:
            file_content: PDF file content as bytes

        Returns:
            Extracted text content
        """
        import io
        text_content = []
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text.strip():
                text_content.append(f"--- Page {page_num} ---\n{page_text}")
        return "\n\n".join(text_content)

    def _extract_docx_text_from_bytes(self, file_content: bytes) -> str:
        """Extract text from DOCX file content.

        Args:
            file_content: DOCX file content as bytes

        Returns:
            Extracted text content
        """
        import io
        docx_file = io.BytesIO(file_content)
        doc = DocxDocument(docx_file)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    def _extract_xlsx_text_from_bytes(self, file_content: bytes) -> str:
        """Extract text from XLSX file content.

        Args:
            file_content: XLSX file content as bytes

        Returns:
            Extracted text content
        """
        import io
        xlsx_file = io.BytesIO(file_content)
        workbook = openpyxl.load_workbook(xlsx_file)
        text_content = []

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_text = [f"--- Sheet: {sheet_name} ---"]

            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    sheet_text.append(row_text)

            text_content.extend(sheet_text)

        return "\n".join(text_content)

    def query_documents(
        self,
        query: str,
        top_k: int = 5,
        similarity_threshold: float = 0.3,
        document_filter: dict[str, Any] | None = None,
    ) -> list[ProcessedDocument]:
        """Query documents using semantic search.

        Args:
            query: Search query
            top_k: Number of top results to return
            similarity_threshold: Minimum similarity threshold
            document_filter: Optional filter for documents

        Returns:
            List of relevant processed documents
        """
        try:
            # Generate query embedding
            formatted_query = self._format_text_for_embedding(query, is_query=True)
            query_embedding = self.embedding_model.encode(formatted_query).tolist()

            # Perform similarity search
            search_results = self.vector_repository.similarity_search(
                query_embedding=query_embedding,
                limit=top_k * 2,  # Get more results for filtering
                document_filter=document_filter
            )

            # Filter by similarity threshold and convert to ProcessedDocuments
            processed_docs = []
            seen_documents = set()

            for result in search_results:
                if result.score >= similarity_threshold and result.document_id not in seen_documents:
                    # Create a simplified ProcessedDocument from the search result
                    # Instead of fetching all chunks, use just this result
                    processed_doc = ProcessedDocument(
                        id=result.document_id,
                        original_document=Document(
                            id=result.document_id,
                            name=result.metadata.get("document_name", "Unknown"),
                            file_path="",  # Not stored in chunks
                            file_type=result.metadata.get("file_type", "unknown"),
                            directory_path=""
                        ),
                        content=result.content,  # Use the chunk content
                        chunks=[result.content],  # Single chunk
                        processed_at=datetime.now(UTC),
                        metadata=result.metadata
                    )

                    processed_docs.append(processed_doc)
                    seen_documents.add(result.document_id)

            return processed_docs[:top_k]

        except Exception as e:
            print(f"Error during document query: {e}")
            return []

    async def search_documents(
        self,
        query: str,
        top_k: int = 5,
        similarity_threshold: float = 0.3,
        document_filter: dict[str, Any] | None = None,
    ) -> list[ProcessedDocument]:
        """Async wrapper for query_documents to maintain compatibility with chat service.

        Args:
            query: Search query
            top_k: Number of top results to return
            similarity_threshold: Minimum similarity threshold
            document_filter: Optional filter for documents

        Returns:
            List of relevant processed documents
        """
        return self.query_documents(query, top_k, similarity_threshold, document_filter)

    def get_collection_stats(self) -> dict[str, Any]:
        """Get statistics about the document collection.

        Returns:
            Dictionary with collection statistics
        """
        try:
            total_chunks = self.vector_repository.get_collection_count()
        except Exception as e:
            print(f"Warning: Could not get collection count: {e}")
            total_chunks = 0

        try:
            processed_documents = len(self.manifest._manifest.get("documents", {}))
        except Exception as e:
            print(f"Warning: Could not get manifest documents: {e}")
            processed_documents = 0

        return {
            "total_chunks": total_chunks,
            "processed_documents": processed_documents
        }

    def get_manifest_stats(self) -> dict[str, Any]:
        """Get manifest statistics for the admin dashboard.

        Returns:
            Dictionary with manifest statistics
        """
        manifest_data = self.manifest._manifest.get("documents", {})
        total_chunks = sum(doc.get("chunk_count", 0) for doc in manifest_data.values())

        return {
            "total_documents": len(manifest_data),
            "total_chunks": total_chunks,
            "embedding_model": self.embedding_model_name
        }

    def delete_document(self, document_id: str) -> bool:
        """Delete document and all its chunks.

        Args:
            document_id: ID of document to delete

        Returns:
            True if deletion was successful
        """
        try:
            deleted_count = self.vector_repository.delete_document_chunks(document_id)
            if deleted_count > 0:
                self.manifest.remove_document(document_id)
                return True
            return False
        except Exception as e:
            print(f"Error deleting document {document_id}: {e}")
            return False

    async def process_documents(self, documents: list[Document]) -> list[ProcessedDocument]:
        """Process documents into searchable format.

        Args:
            documents: List of raw documents

        Returns:
            List of processed documents
        """
        processed_docs = []
        for document in documents:
            try:
                processed_doc = self.add_document(document)
                processed_docs.append(processed_doc)
            except Exception as e:
                print(f"Error processing document {document.name}: {e}")
                continue
        return processed_docs

    async def process_changed_documents(self, documents: list[Document]) -> list[ProcessedDocument]:
        """Process only changed documents based on manifest tracking.

        Args:
            documents: List of all documents to check for changes

        Returns:
            List of processed documents that were changed
        """
        changed_docs = self.manifest.get_changed_documents(documents)
        if not changed_docs:
            print("No changed documents found")
            return []

        print(f"Processing {len(changed_docs)} changed documents")
        return await self.process_documents(changed_docs)

    async def index_documents(self, processed_docs: list[ProcessedDocument]) -> None:
        """Index processed documents for fast retrieval.

        Args:
            processed_docs: List of processed documents to index
        """
        # Documents are automatically indexed when added via add_document
        # This method is here for interface compatibility
        print(f"Documents are automatically indexed in MongoDB. {len(processed_docs)} docs ready.")
        pass
