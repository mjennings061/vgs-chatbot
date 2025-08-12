"""Document processor service implementation with RAG pipeline."""

import hashlib
import uuid
from datetime import UTC, datetime
from io import BytesIO
from typing import Any

import chromadb
import openpyxl
import tiktoken
from docx import Document as DocxDocument
from pypdf import PdfReader
from sentence_transformers import SentenceTransformer

from vgs_chatbot.interfaces.document_processor_interface import (
    DocumentProcessorInterface,
)
from vgs_chatbot.models.document import Document, ProcessedDocument
from vgs_chatbot.utils.manifest import DocumentManifest
from vgs_chatbot.utils.retrieval_enhancer import RetrievalEnhancer


class RAGDocumentProcessor(DocumentProcessorInterface):
    """RAG-based document processor implementation."""

    def __init__(
        self,
        embedding_model: str = "BAAI/bge-small-en-v1.5",
        persist_directory: str | None = None,
    ) -> None:
        """Initialize document processor.

        Args:
            embedding_model: Name of sentence transformer model for embeddings
            persist_directory: Optional path for persistent Chroma storage
        """
        self.embedding_model_name = embedding_model

        # Try preferred model with fallback
        try:
            self.embedding_model = SentenceTransformer(embedding_model)
            print(f"Loaded embedding model: {embedding_model}")
        except Exception as e:
            print(
                f"Failed to load {embedding_model}, falling back to multi-qa-MiniLM-L6-cos-v1: {e}"
            )
            self.embedding_model_name = "multi-qa-MiniLM-L6-cos-v1"
            self.embedding_model = SentenceTransformer("multi-qa-MiniLM-L6-cos-v1")

        self.persist_directory = persist_directory

        if persist_directory:
            # Use persistent client to retain embeddings across Streamlit reruns
            self.chroma_client = chromadb.PersistentClient(path=persist_directory)
            # Initialize manifest in same directory
            manifest_path = f"{persist_directory}/manifest.json"
        else:
            self.chroma_client = chromadb.Client()
            manifest_path = "data/vectors/manifest.json"

        self.manifest = DocumentManifest(manifest_path)
        self.retrieval_enhancer = RetrievalEnhancer()

        # Try to get existing collection; create if not present with explicit cosine distance
        try:
            self.collection = self.chroma_client.get_collection(name="documents")
        except Exception:
            self.collection = self.chroma_client.create_collection(
                name="documents", metadata={"hnsw:space": "cosine"}
            )
        # Section-based chunking parameters
        self.chunk_target_tokens = 450  # Target tokens per chunk
        self.chunk_max_tokens = 1000  # Maximum tokens for large sections

        # Initialize tiktoken encoder for accurate token counting
        self.tokenizer = tiktoken.get_encoding(
            "cl100k_base"
        )  # GPT-4 compatible encoding

        # Determine model-specific formatting needs
        self._requires_prefixes = self._model_requires_prefixes(
            self.embedding_model_name
        )
        if self._requires_prefixes:
            print(f"Model {self.embedding_model_name} requires query/passage prefixes")

        # Deduplication tracking
        self._chunk_hashes: set[str] = set()

        # BM25 index for hybrid retrieval (initialized when documents are indexed)
        self._bm25_index = None
        self._bm25_corpus = []
        self._bm25_ids = []

    def _token_count(self, text: str) -> int:
        """Count tokens in text using tiktoken encoder.

        Args:
            text: Text to count tokens for

        Returns:
            Number of tokens in text
        """
        return len(self.tokenizer.encode(text))

    def _normalize_text_for_dedup(self, text: str) -> str:
        """Normalize text for deduplication by removing formatting differences.

        Args:
            text: Text to normalize

        Returns:
            Normalized text for hashing
        """
        # Convert to lowercase and normalize whitespace
        normalized = " ".join(text.lower().split())
        # Remove common formatting artifacts
        normalized = normalized.replace("•", "").replace("-", "").replace("*", "")
        # Remove page markers and document artifacts
        import re

        normalized = re.sub(r"--- page \d+ ---", "", normalized)
        normalized = re.sub(r"=== .* ===", "", normalized)
        return normalized.strip()

    def _get_text_hash(self, text: str) -> str:
        """Generate hash for text deduplication.

        Args:
            text: Text to hash

        Returns:
            SHA256 hash of normalized text
        """
        normalized = self._normalize_text_for_dedup(text)
        return hashlib.sha256(normalized.encode("utf-8")).hexdigest()

    def _is_duplicate_chunk(self, chunk: str) -> bool:
        """Check if chunk is a duplicate based on content hash.

        Args:
            chunk: Text chunk to check

        Returns:
            True if chunk is a duplicate
        """
        chunk_hash = self._get_text_hash(chunk)
        if chunk_hash in self._chunk_hashes:
            return True
        self._chunk_hashes.add(chunk_hash)
        return False

    def _extract_header_path(self, lines: list[str], current_line_idx: int) -> str:
        """Extract header path context for a given line position.

        Args:
            lines: All lines from the document
            current_line_idx: Current line index to extract context for

        Returns:
            Header path string (e.g. "DHO 2305 > Annex E > Weather Limitations")
        """
        headers = []
        # Look backwards from current position to find recent headings
        for i in range(max(0, current_line_idx - 50), current_line_idx):
            line = lines[i].strip()
            if line:
                # Check for enhanced structure headings
                if line.startswith("===") and line.endswith("==="):
                    header = line.replace("===", "").strip()
                    if header:
                        headers.append(header)
                # Check for markdown-style headings
                elif line.startswith("#"):
                    header = line.lstrip("#").strip()
                    if header:
                        headers.append(header)
                # Check for document reference patterns (DHO, Annex, etc.)
                elif any(
                    pattern in line.upper()
                    for pattern in ["DHO ", "ANNEX ", "APPENDIX ", "SECTION "]
                ):
                    if len(line) < 100:  # Avoid long paragraphs
                        headers.append(line)

        # Return the last 3 headers to avoid overly long paths
        return " > ".join(headers[-3:]) if headers else ""

    async def process_documents(
        self, documents: list[Document]
    ) -> list[ProcessedDocument]:
        """Process documents into searchable format.

        Args:
            documents: List of raw documents

        Returns:
            List of processed documents
        """
        processed_docs = []

        for document in documents:
            try:
                processed_doc = await self._process_single_document(document)
                processed_docs.append(processed_doc)
            except Exception as e:
                # Log error but continue with other documents
                print(f"Error processing document {document.name}: {e}")
                continue

        return processed_docs

    async def _process_single_document(self, document: Document) -> ProcessedDocument:
        """Process a single document.

        Args:
            document: Document to process

        Returns:
            Processed document
        """
        # Extract text content based on file type
        content = await self._extract_text_content(document)

        # Split into chunks with safety validation
        chunks = self._split_text_into_chunks(content)

        # Safety validation: filter out empty and invalid chunks
        validated_chunks = self._validate_and_filter_chunks(chunks, document.name)

        if not validated_chunks:
            print(f"Warning: No valid chunks generated for {document.name}")
            # Return empty processed document
            return ProcessedDocument(
                id=str(uuid.uuid4()),
                original_document=document,
                content=content,
                chunks=[],
                embeddings=None,
                metadata=self._extract_metadata(content, document.name),
                processed_at=datetime.now(UTC),
            )

        # Generate embeddings for validated chunks with normalization and formatting
        try:
            # Format chunks for embedding (add passage prefixes if needed)
            formatted_chunks = [
                self._format_for_embedding(chunk, is_query=False)
                for chunk in validated_chunks
            ]

            raw_embeddings: Any = self.embedding_model.encode(
                formatted_chunks, convert_to_numpy=True, normalize_embeddings=True
            )  # type: ignore[no-untyped-call]
            if hasattr(raw_embeddings, "tolist"):
                embeddings = raw_embeddings.tolist()  # type: ignore[assignment]
            else:  # Fallback (already python list)
                embeddings = raw_embeddings  # type: ignore[assignment]
            print(
                f"Generated {len(embeddings)} normalized embeddings for {len(validated_chunks)} validated chunks"
            )
            if self._requires_prefixes:
                print(f"Applied passage prefixes for {self.embedding_model_name} model")
        except Exception as e:
            print(f"Error generating embeddings: {e}")
            embeddings = None

        # Generate document-level summary embedding
        doc_summary_embedding = None
        try:
            doc_summary = self._generate_document_summary(content, document.name)
            if doc_summary:
                formatted_summary = self._format_for_embedding(
                    doc_summary, is_query=False
                )
                summary_emb: Any = self.embedding_model.encode(
                    [formatted_summary],
                    convert_to_numpy=True,
                    normalize_embeddings=True,
                )  # type: ignore[no-untyped-call]
                if hasattr(summary_emb, "tolist"):
                    doc_summary_embedding = summary_emb.tolist()[0]  # type: ignore[assignment]
                else:
                    doc_summary_embedding = summary_emb[0]  # type: ignore[assignment]
                print(f"Generated document-level summary embedding for {document.name}")
        except Exception as e:
            print(f"Error generating document summary embedding: {e}")
            doc_summary_embedding = None

        # Extract metadata for better context
        metadata = self._extract_metadata(content, document.name)

        processed_doc = ProcessedDocument(
            id=str(uuid.uuid4()),
            original_document=document,
            content=content,
            chunks=validated_chunks,
            embeddings=embeddings,
            metadata=metadata,
            processed_at=datetime.now(UTC),
        )

        # Store document summary embedding separately
        if doc_summary_embedding is not None:
            processed_doc._doc_summary_embedding = doc_summary_embedding  # type: ignore
            processed_doc._doc_summary = doc_summary  # type: ignore

        return processed_doc

    async def _extract_text_content(self, document: Document) -> str:
        """Extract text content from document based on file type.

        Args:
            document: Document to extract text from

        Returns:
            Extracted text content
        """
        from pathlib import Path

        try:
            file_path = Path(document.file_path)
            if not file_path.exists():
                return f"Document: {document.name}\nError: File not found at {document.file_path}"

            # Read file content as bytes
            file_content = file_path.read_bytes()

            if (
                document.file_type == "application/pdf"
                or file_path.suffix.lower() == ".pdf"
            ):
                return await self._extract_pdf_text(file_content)
            elif document.file_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ] or file_path.suffix.lower() in [".docx", ".doc"]:
                return await self._extract_docx_text(file_content)
            elif document.file_type in [
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "application/vnd.ms-excel",
            ] or file_path.suffix.lower() in [".xlsx", ".xls"]:
                return await self._extract_excel_text(file_content)
            elif file_path.suffix.lower() in [".txt", ".md"]:
                # Handle text files directly
                return file_content.decode("utf-8", errors="replace")
            else:
                # Try to decode as text
                try:
                    return file_content.decode("utf-8", errors="replace")
                except Exception:
                    return f"Document: {document.name}\nBinary file content cannot be extracted as text"

        except Exception as e:
            return f"Document: {document.name}\nError extracting content: {str(e)}"

    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF content with improved structure detection.

        Args:
            content: PDF file content as bytes

        Returns:
            Extracted text with page markers and structure
        """
        if not content:
            return "PDF content placeholder"

        try:
            reader = PdfReader(BytesIO(content))
            text = ""
            total_pages = len(reader.pages)

            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        # Add page marker with more context
                        text += f"\n--- PAGE {page_num} of {total_pages} ---\n"

                        # Try to preserve more structure
                        enhanced_text = self._enhance_pdf_text_structure(page_text)
                        text += enhanced_text + "\n"

                except Exception as e:
                    # Log individual page errors but continue processing
                    print(f"Warning: Error extracting page {page_num}: {e}")
                    text += f"\n--- PAGE {page_num} (EXTRACTION ERROR) ---\n"
                    continue

            # Clean up text - remove excessive whitespace but preserve structure
            text = self._clean_extracted_text(text)
            return text if text else "No text content found in PDF"

        except Exception as e:
            return f"Error extracting PDF text: {str(e)}"

    def _enhance_pdf_text_structure(self, page_text: str) -> str:
        """Enhance PDF text structure by identifying headings, lists, and paragraphs.

        Args:
            page_text: Raw text from PDF page

        Returns:
            Enhanced text with better structure
        """
        lines = page_text.split("\n")
        enhanced_lines = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Identify potential headings (all caps, short lines, etc.)
            if len(line) < 80 and (
                line.isupper()
                or any(
                    word in line.upper()
                    for word in ["CHAPTER", "SECTION", "PART", "APPENDIX"]
                )
                or line.endswith(":")
            ):
                enhanced_lines.append(f"\n=== {line} ===\n")

            # Identify bullet points and lists
            elif any(line.startswith(marker) for marker in ["•", "-", "*", "○"]):
                enhanced_lines.append(f"  {line}")

            # Identify numbered lists
            elif len(line) > 3 and line[:3].replace(".", "").replace(")", "").isdigit():
                enhanced_lines.append(f"  {line}")

            # Regular text
            else:
                enhanced_lines.append(line)

        return "\n".join(enhanced_lines)

    def _clean_extracted_text(self, text: str) -> str:
        """Clean extracted text while preserving important structure.

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text with preserved structure
        """
        lines = text.split("\n")
        cleaned_lines = []

        for line in lines:
            line = line.strip()
            if line:
                cleaned_lines.append(line)
            elif cleaned_lines and not cleaned_lines[-1] == "":
                # Preserve paragraph breaks
                cleaned_lines.append("")

        # Remove excessive empty lines (more than 2 consecutive)
        result_lines = []
        empty_count = 0

        for line in cleaned_lines:
            if line == "":
                empty_count += 1
                if empty_count <= 2:
                    result_lines.append(line)
            else:
                empty_count = 0
                result_lines.append(line)

        return "\n".join(result_lines)

    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX content with structure preservation.

        Args:
            content: DOCX file content as bytes

        Returns:
            Extracted text with structure and tables
        """
        if not content:
            return "DOCX content placeholder"

        try:
            doc = DocxDocument(BytesIO(content))
            text = ""

            # Extract paragraphs with style information
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    # Check if paragraph is a heading
                    style_name = (
                        getattr(getattr(paragraph, "style", None), "name", "") or ""
                    )
                    if style_name.startswith("Heading"):
                        text += f"\n## {para_text}\n"
                    else:
                        text += para_text + "\n"

            # Extract tables
            for table_num, table in enumerate(doc.tables, 1):
                text += f"\n--- TABLE {table_num} ---\n"
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_cells.append(cell_text)
                    if row_cells:
                        text += " | ".join(row_cells) + "\n"
                text += "\n"

            return text if text.strip() else "No text content found in DOCX"

        except Exception as e:
            return f"Error extracting DOCX text: {str(e)}"

    async def _extract_excel_text(self, content: bytes) -> str:
        """Extract text from Excel content with improved structure handling.

        Args:
            content: Excel file content as bytes

        Returns:
            Extracted text with sheet and row structure
        """
        if not content:
            return "Excel content placeholder"

        try:
            workbook = openpyxl.load_workbook(BytesIO(content), data_only=True)
            text = ""

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\n--- SHEET: {sheet_name} ---\n"

                # Get data range to avoid empty rows
                if sheet.max_row > 0:
                    # Extract header row if it exists
                    first_row = list(
                        sheet.iter_rows(min_row=1, max_row=1, values_only=True)
                    )[0]
                    if first_row and any(cell is not None for cell in first_row):
                        headers = [
                            str(cell) if cell is not None else "" for cell in first_row
                        ]
                        text += "Headers: " + " | ".join(headers) + "\n"
                        start_row = 2
                    else:
                        start_row = 1

                    # Extract data rows
                    row_count = 0
                    for row in sheet.iter_rows(min_row=start_row, values_only=True):
                        if row and any(cell is not None for cell in row):
                            row_text = " | ".join(
                                str(cell) if cell is not None else "" for cell in row
                            )
                            if row_text.strip():
                                text += row_text + "\n"
                                row_count += 1
                                # Limit rows to prevent excessive content
                                if row_count > 100:
                                    text += "[... truncated after 100 rows ...]\n"
                                    break

            return text if text.strip() else "No data found in Excel file"

        except Exception as e:
            return f"Error extracting Excel text: {str(e)}"

    def _split_text_into_chunks(self, text: str) -> list[str]:
        """Split text into token-based chunks with overlap and header context preservation.

        Args:
            text: Text to split

        Returns:
            List of text chunks with header context and proper token sizing
        """
        # Special handling for weather limitations table - preserve as single chunk
        weather_table_chunk = self._extract_weather_limitations_table(text)
        if weather_table_chunk:
            # Remove the table from the original text to avoid duplication
            text_without_table = text.replace(
                weather_table_chunk, "\n[WEATHER_TABLE_REMOVED]\n"
            )

            # Process the rest of the document normally
            other_chunks = self._split_text_standard(text_without_table)

            # Insert the weather table as the first chunk (highest priority)
            return [weather_table_chunk] + other_chunks

        # Standard processing with section-based chunking
        return self._split_text_standard(text)

    def _extract_weather_limitations_table(self, text: str) -> str | None:
        """Extract the complete weather limitations table as a single chunk."""
        lines = text.split("\n")

        # Find weather table boundaries
        start_idx = -1
        end_idx = -1

        for i, line in enumerate(lines):
            line_lower = line.lower().strip()

            # Look for table start
            if "weather limitations table" in line_lower:
                start_idx = i

                # Look for the end - find where wind/crosswind data ends
                for j in range(i + 1, min(i + 25, len(lines))):
                    next_line = lines[j].strip()

                    # Table continues while we have wind/crosswind data
                    if (
                        (
                            "crosswind" in next_line.lower()
                            and "kts" in next_line.lower()
                        )
                        or ("maximum wind speed" in next_line.lower())
                        or (
                            (
                                "20kts" in next_line
                                or "25kts" in next_line
                                or "30kts" in next_line
                            )
                            and len(next_line.split()) < 20
                        )
                    ):  # Short lines with wind data
                        end_idx = j
                    # Stop at major section break
                    elif (
                        (
                            next_line.startswith("===")
                            and "weather" not in next_line.lower()
                            and "cfs" not in next_line.lower()
                        )
                        or next_line.startswith("DHO ")
                        or "minimum cloud base" in next_line.lower()
                    ):
                        break

                break

        # Extract the complete table if found
        if start_idx >= 0 and end_idx > start_idx:
            # Include some context before the table
            context_start = max(0, start_idx - 2)
            table_lines = lines[context_start : end_idx + 1]
            return "\n".join(table_lines).strip()

        return None

    def _split_text_token_based(self, text: str) -> list[str]:
        """Split text into token-based chunks with overlap and header awareness.

        Args:
            text: Text to split

        Returns:
            List of text chunks with header context prefixes
        """
        lines = text.split("\n")

        # First, identify blocks (paragraphs, sections, tables)
        blocks = []
        current_block: list[str] = []

        for i, line in enumerate(lines):
            line_stripped = line.strip()

            # Skip removed table markers
            if line_stripped == "[WEATHER_TABLE_REMOVED]":
                continue

            # Detect block boundaries (empty lines, page markers, section headers)
            is_boundary = (
                not line_stripped  # Empty line
                or line_stripped.startswith("--- PAGE")  # Page marker
                or (
                    line_stripped.startswith("===") and line_stripped.endswith("===")
                )  # Section header
                or line_stripped.startswith("#")  # Markdown header
            )

            if is_boundary and current_block:
                # Save current block
                blocks.append(
                    {
                        "lines": current_block.copy(),
                        "start_line": i - len(current_block),
                        "content": "\n".join(current_block),
                    }
                )
                current_block = []

            if line_stripped:  # Don't add empty lines to blocks
                current_block.append(line)

        # Add final block
        if current_block:
            blocks.append(
                {
                    "lines": current_block.copy(),
                    "start_line": len(lines) - len(current_block),
                    "content": "\n".join(current_block),
                }
            )

        # Now chunk these blocks with token-based sizing and overlap
        return self._chunk_blocks_with_overlap(blocks, lines)

    def _chunk_blocks_with_overlap(
        self, blocks: list[dict[str, Any]], all_lines: list[str]
    ) -> list[str]:
        """Chunk blocks with token-based sizing and overlap.

        Args:
            blocks: List of content blocks with metadata
            all_lines: All document lines for header context

        Returns:
            List of chunks with header context
        """
        chunks = []
        current_chunk_blocks = []
        current_tokens = 0

        for block in blocks:
            block_content = block["content"]
            block_tokens = self._token_count(block_content)

            # Check if adding this block would exceed our target size
            if current_tokens + block_tokens <= self.chunk_target_tokens:
                current_chunk_blocks.append(block)
                current_tokens += block_tokens
            else:
                # Save current chunk if it has content
                if current_chunk_blocks:
                    chunk_text = self._create_chunk_with_context(
                        current_chunk_blocks, all_lines
                    )
                    chunks.append(chunk_text)

                # Handle oversized blocks
                if block_tokens > self.chunk_target_tokens:
                    # Split large block at sentence/line boundaries
                    sub_chunks = self._split_large_block(block, all_lines)
                    chunks.extend(sub_chunks)
                    current_chunk_blocks = []
                    current_tokens = 0
                else:
                    # Start new chunk with this block
                    current_chunk_blocks = [block]
                    current_tokens = block_tokens

        # Add final chunk
        if current_chunk_blocks:
            chunk_text = self._create_chunk_with_context(
                current_chunk_blocks, all_lines
            )
            chunks.append(chunk_text)

        # Return chunks without overlap to preserve section boundaries
        return chunks if chunks else []

    def _create_chunk_with_context(
        self, chunk_blocks: list[dict[str, Any]], all_lines: list[str]
    ) -> str:
        """Create a chunk with header context prefix.

        Args:
            chunk_blocks: Blocks to include in chunk
            all_lines: All document lines for header context

        Returns:
            Chunk text with header context prefix
        """
        if not chunk_blocks:
            return ""

        # Get header context from the first block's position
        first_block = chunk_blocks[0]
        header_path = self._extract_header_path(all_lines, first_block["start_line"])

        # Combine block content
        chunk_content = "\n\n".join(block["content"] for block in chunk_blocks)

        # Add header context if available
        if header_path:
            return f"{header_path}\n\n{chunk_content}"
        else:
            return chunk_content

    def _split_large_block(
        self, block: dict[str, Any], all_lines: list[str]
    ) -> list[str]:
        """Split a large block that exceeds token limits.

        Args:
            block: Block to split
            all_lines: All document lines for header context

        Returns:
            List of sub-chunks
        """
        content = block["content"]
        lines = content.split("\n")
        sub_chunks = []
        current_lines = []
        current_tokens = 0

        # Get header context for this block
        header_path = self._extract_header_path(all_lines, block["start_line"])
        header_prefix = f"{header_path}\n\n" if header_path else ""
        header_tokens = self._token_count(header_prefix)

        for line in lines:
            line_tokens = self._token_count(line)

            if current_tokens + line_tokens + header_tokens <= self.chunk_target_tokens:
                current_lines.append(line)
                current_tokens += line_tokens
            else:
                # Save current sub-chunk
                if current_lines:
                    sub_chunk = header_prefix + "\n".join(current_lines)
                    sub_chunks.append(sub_chunk)

                # Start new sub-chunk
                current_lines = [line]
                current_tokens = line_tokens

        # Add final sub-chunk
        if current_lines:
            sub_chunk = header_prefix + "\n".join(current_lines)
            sub_chunks.append(sub_chunk)

        return sub_chunks

    def _add_chunk_overlap(self, chunks: list[str]) -> list[str]:
        """Add overlap between adjacent chunks.

        Args:
            chunks: List of chunks to add overlap to

        Returns:
            List of chunks with overlap added
        """
        if len(chunks) <= 1:
            return chunks

        overlapped_chunks = [chunks[0]]  # First chunk unchanged

        for i in range(1, len(chunks)):
            prev_chunk = chunks[i - 1]
            current_chunk = chunks[i]

            # Extract tail tokens from previous chunk for overlap
            prev_tokens = self.tokenizer.encode(prev_chunk)
            overlap_tokens = (
                prev_tokens[-self.chunk_overlap_tokens :]
                if len(prev_tokens) > self.chunk_overlap_tokens
                else prev_tokens
            )
            overlap_text = self.tokenizer.decode(overlap_tokens)

            # Combine overlap with current chunk
            overlapped_chunk = f"{overlap_text}\n\n{current_chunk}"
            overlapped_chunks.append(overlapped_chunk)

        return overlapped_chunks

    def _identify_aviation_sections(self, text: str) -> list[str]:
        """Identify sections in aviation documents using aviation-specific patterns.

        Args:
            text: Document text to analyze

        Returns:
            List of identified sections
        """
        lines = text.split("\n")
        sections = []
        current_section = ""

        for i, line in enumerate(lines):
            line_stripped = line.strip()

            # Skip removed table markers
            if line_stripped == "[WEATHER_TABLE_REMOVED]":
                continue

            # Check if this line is a section boundary
            is_section_boundary = self._is_aviation_section_boundary(
                line_stripped, i, lines
            )

            if is_section_boundary and current_section.strip():
                # Save the current section before starting a new one
                sections.append(current_section.strip())
                current_section = line + "\n"
            else:
                current_section += line + "\n"

        # Add the final section
        if current_section.strip():
            sections.append(current_section.strip())

        return sections if sections else []

    def _is_aviation_section_boundary(
        self, line: str, line_idx: int, all_lines: list[str]
    ) -> bool:
        """Determine if a line represents an aviation section boundary.

        Args:
            line: Current line to check
            line_idx: Index of current line
            all_lines: All lines in the document

        Returns:
            True if this line starts a new section
        """
        if not line:
            return False

        # Page markers
        if line.startswith("--- PAGE"):
            return True

        # Enhanced structure headings (=== ... ===)
        if line.startswith("===") and line.endswith("==="):
            return True

        # Markdown-style headings
        if line.startswith("#"):
            return True

        # Aviation document patterns
        aviation_section_patterns = [
            # DHO/Document references
            r"^DHO\s+\d+",
            r"^ANNEX\s+[A-Z]",
            r"^APPENDIX\s+[A-Z\d]",
            r"^SECTION\s+\d+",
            r"^CHAPTER\s+\d+",
            r"^PART\s+[A-Z\d]",
            # Aviation operational sections
            r"^WEATHER\s+LIMITATIONS?",
            r"^OPERATIONAL\s+LIMITS?",
            r"^EMERGENCY\s+PROCEDURES?",
            r"^FLIGHT\s+OPERATIONS?",
            r"^SAFETY\s+PROCEDURES?",
            r"^MAINTENANCE\s+PROCEDURES?",
            # Numbered procedures and lists (start of major sections)
            r"^\d+\.\s+[A-Z].*PROCEDURES?",
            r"^\d+\.\s+[A-Z].*OPERATIONS?",
            r"^\d+\.\s+[A-Z].*LIMITS?",
        ]

        import re

        for pattern in aviation_section_patterns:
            if re.match(pattern, line.upper()):
                return True

        # Check for all-caps section titles (common in aviation docs)
        if (
            line.isupper()
            and len(line) > 10
            and len(line) < 100
            and any(
                word in line
                for word in [
                    "OPERATIONS",
                    "PROCEDURES",
                    "LIMITS",
                    "WEATHER",
                    "SAFETY",
                    "EMERGENCY",
                ]
            )
        ):

            # Make sure this isn't just a sentence in the middle of a paragraph
            # Check if previous line was empty or is start of document
            if line_idx == 0 or (line_idx > 0 and not all_lines[line_idx - 1].strip()):
                return True

        return False

    def _split_text_standard(self, text: str) -> list[str]:
        """Section-based text splitting that preserves document structure."""
        chunks = []

        # Primary: Identify sections based on document structure
        sections = self._identify_document_sections(text)

        # Fallback 1: Split by page markers if no clear sections
        if len(sections) <= 1:
            sections = self._split_by_page_markers(text)

        # Fallback 2: Split by paragraph breaks
        if len(sections) <= 1:
            sections = [s.strip() for s in text.split("\n\n") if s.strip()]

        # Final fallback: Use entire text
        if len(sections) <= 1:
            sections = [text.strip()]

        # Process sections with size management
        current_chunk = ""
        for section in sections:
            if not section or section == "[WEATHER_TABLE_REMOVED]":
                continue

            section_tokens = self._token_count(section)

            # If section is small enough, try to combine with previous chunks
            if current_chunk and self._token_count(current_chunk + "\n\n" + section) <= self.chunk_target_tokens:
                current_chunk = (current_chunk + "\n\n" + section).strip()
            else:
                # Save current chunk if it has content
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())

                # Handle oversized sections
                if section_tokens > self.chunk_max_tokens:
                    # Split oversized sections while trying to preserve subsection boundaries
                    sub_chunks = self._split_oversized_section(section)
                    chunks.extend(sub_chunks)
                    current_chunk = ""
                else:
                    # Start new chunk with this section
                    current_chunk = section.strip()

        # Add final chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        return chunks if chunks else [text.strip()]

    def _identify_document_sections(self, text: str) -> list[str]:
        """Identify document sections based on structural patterns.
        
        Args:
            text: Document text to analyze
            
        Returns:
            List of identified sections
        """
        lines = text.split("\n")
        sections = []
        current_section = ""

        for i, line in enumerate(lines):
            line_stripped = line.strip()

            # Skip removed table markers
            if line_stripped == "[WEATHER_TABLE_REMOVED]":
                continue

            # Check if this line is a section boundary
            is_section_boundary = self._is_section_boundary(line_stripped, i, lines)

            if is_section_boundary and current_section.strip():
                # Save the current section before starting a new one
                sections.append(current_section.strip())
                current_section = line + "\n"
            else:
                current_section += line + "\n"

        # Add the final section
        if current_section.strip():
            sections.append(current_section.strip())

        return sections if sections else []

    def _is_section_boundary(self, line: str, line_idx: int, all_lines: list[str]) -> bool:
        """Determine if a line represents a section boundary.
        
        Args:
            line: Current line to check
            line_idx: Index of current line
            all_lines: All lines in the document
            
        Returns:
            True if this line starts a new section
        """
        if not line:
            return False

        # Page markers
        if line.startswith("--- PAGE"):
            return True

        # Enhanced structure headings (=== ... ===)
        if line.startswith("===") and line.endswith("==="):
            return True

        # Markdown-style headings
        if line.startswith("#"):
            return True

        # Document reference patterns (general)
        import re
        section_patterns = [
            r"^DHO\s+\d+",
            r"^ANNEX\s+[A-Z]",
            r"^APPENDIX\s+[A-Z\d]",
            r"^SECTION\s+\d+",
            r"^CHAPTER\s+\d+",
            r"^PART\s+[A-Z\d]",
        ]

        for pattern in section_patterns:
            if re.match(pattern, line.upper()):
                return True

        # All-caps titles that appear to be headings
        if (line.isupper() and
            len(line) > 10 and
            len(line) < 100 and
            # Make sure this isn't just a sentence in the middle of a paragraph
            (line_idx == 0 or not all_lines[line_idx - 1].strip())):
            return True

        return False

    def _split_by_page_markers(self, text: str) -> list[str]:
        """Split text by page markers as fallback method.
        
        Args:
            text: Text to split
            
        Returns:
            List of page-based sections
        """
        sections = []
        current_section = ""

        for line in text.split("\n"):
            if line.startswith("--- PAGE"):
                if current_section.strip():
                    sections.append(current_section.strip())
                current_section = line + "\n"
            else:
                current_section += line + "\n"

        # Add the last section
        if current_section.strip():
            sections.append(current_section.strip())

        return sections

    def _split_oversized_section(self, section: str) -> list[str]:
        """Split an oversized section while preserving subsection boundaries.
        
        Args:
            section: Section text to split
            
        Returns:
            List of sub-chunks
        """
        # Try to split by subsection markers first
        lines = section.split("\n")
        sub_chunks = []
        current_chunk = ""

        for line in lines:
            line_stripped = line.strip()

            # Check for subsection boundaries
            is_subsection = (
                line_stripped.startswith("===") or
                line_stripped.startswith("#") or
                (line_stripped.isupper() and len(line_stripped) > 5 and len(line_stripped) < 80)
            )

            if is_subsection and current_chunk.strip() and self._token_count(current_chunk) > 200:
                # Save current chunk if it's substantial
                sub_chunks.append(current_chunk.strip())
                current_chunk = line + "\n"
            else:
                current_chunk += line + "\n"

            # Check if current chunk is getting too large
            if self._token_count(current_chunk) > self.chunk_target_tokens:
                if current_chunk.strip():
                    sub_chunks.append(current_chunk.strip())
                current_chunk = ""

        # Add final chunk
        if current_chunk.strip():
            sub_chunks.append(current_chunk.strip())

        return sub_chunks if sub_chunks else [section]

    def _extract_metadata(self, content: str, document_name: str) -> dict[str, Any]:
        """Extract metadata from document content.

        Args:
            content: Document text content
            document_name: Name of the document

        Returns:
            Dictionary containing extracted metadata
        """
        metadata = {
            "document_name": document_name,
            "total_pages": 0,
            "sections": [],
            "key_terms": [],
        }

        # Count pages from page markers
        page_markers = content.count("--- PAGE")
        if page_markers > 0:
            metadata["total_pages"] = page_markers

        # Extract section headings (improved heuristic)
        lines = content.split("\n")
        sections = []
        for line in lines:
            line = line.strip()
            # Look for enhanced structure headings
            if line.startswith("===") and line.endswith("==="):
                # Extract content between === markers
                section_title = line.replace("===", "").strip()
                if section_title:
                    sections.append(section_title)
            # Look for traditional section patterns
            elif line.isupper() and len(line) > 5 and len(line) < 100:
                sections.append(line)
            # Look for markdown-style headings
            elif line.startswith("#") and len(line) > 5:
                sections.append(line)
            # Look for numbered sections and chapters
            elif any(
                word in line.upper()
                for word in ["CHAPTER", "SECTION", "PART", "APPENDIX"]
            ):
                sections.append(line)
            # Look for topic-specific sections
            elif any(
                word in line.upper()
                for word in [
                    "WEATHER",
                    "OPERATIONS",
                    "LIMITS",
                    "PROCEDURES",
                    "EMERGENCY",
                    "SAFETY",
                ]
            ):
                if len(line) < 100:  # Avoid capturing long paragraphs
                    sections.append(line)

        metadata["sections"] = sections[:10]  # Limit to first 10 sections

        # Extract aviation-specific key terms that might be relevant
        aviation_terms = [
            "wind limit",
            "wind speed",
            "crosswind",
            "headwind",
            "tailwind",
            "weather",
            "conditions",
            "limits",
            "maximum",
            "minimum",
            "operational",
            "knots",
            "kt",
            "mph",
            "m/s",
            "gusts",
            "turbulence",
            "visibility",
            "gliding",
            "glider",
            "aircraft",
            "operations",
            "safety",
            "emergency",
            "procedure",
            "altitude",
            "approach",
            "landing",
            "takeoff",
            "circuit",
        ]

        found_terms = []
        content_lower = content.lower()
        for term in aviation_terms:
            if term in content_lower:
                found_terms.append(term)

        # Also extract any numeric values that might be limits
        import re

        numeric_patterns = re.findall(
            r"\d+\s*(?:knots|kt|mph|m/s|feet|ft|meters|m)\b", content_lower
        )
        if numeric_patterns:
            # Add unique numeric limits to key terms
            for pattern in set(numeric_patterns[:5]):  # Limit to 5 most common
                found_terms.append(pattern)

        metadata["key_terms"] = found_terms

        return metadata

    async def search_documents(
        self, query: str, top_k: int = 5
    ) -> list[ProcessedDocument]:
        """Search processed documents using hybrid semantic + keyword search.

        Args:
            query: Search query
            top_k: Number of top results to return

        Returns:
            List of relevant processed documents ranked by hybrid score
        """
        # Preprocess query for better search
        enhanced_query = self.retrieval_enhancer.preprocess_query_for_search(query)
        print(f"Enhanced query: '{query}' -> '{enhanced_query[:100]}...'")

        # Format query for embedding (add query prefix if needed)
        formatted_query = self._format_for_embedding(enhanced_query, is_query=True)

        # Generate query embedding with normalization
        try:
            qe: Any = self.embedding_model.encode(
                [formatted_query], convert_to_numpy=True, normalize_embeddings=True
            )  # type: ignore
            if hasattr(qe, "tolist"):
                qe_list = qe.tolist()  # type: ignore
            else:
                qe_list = qe  # type: ignore
            query_embedding = qe_list[0] if qe_list else []  # type: ignore[index]
            print("Generated normalized query embedding for enhanced query")
            if self._requires_prefixes:
                print(f"Applied query prefix for {self.embedding_model_name} model")
        except Exception as e:
            print(f"Error generating query embedding: {e}")
            return []

        # Perform hybrid retrieval: Dense (ChromaDB) + BM25
        collection_count = self.collection.count()
        if collection_count == 0:
            print("Warning: No documents indexed in ChromaDB")
            return []

        # Step 1: Dense retrieval from ChromaDB
        initial_limit = min(top_k * 4, collection_count, 50)
        try:
            dense_results_raw: dict[str, Any] = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=initial_limit,
            )  # type: ignore[assignment]

            # Convert to standardized format
            dense_results = []
            if dense_results_raw.get("ids") and dense_results_raw.get("ids")[0]:
                ids_list = dense_results_raw.get("ids", [[]])[0]
                metadatas_list = dense_results_raw.get("metadatas", [[]])[0]
                documents_list = dense_results_raw.get("documents", [[]])[0]
                distances_list = (
                    dense_results_raw.get("distances", [[]])[0]
                    if dense_results_raw.get("distances")
                    else [0.0] * len(ids_list)
                )

                for i, chunk_id in enumerate(ids_list):
                    try:
                        dense_results.append(
                            {
                                "chunk_id": chunk_id,
                                "content": str(documents_list[i]),
                                "distance": (
                                    distances_list[i]
                                    if i < len(distances_list)
                                    else 1.0
                                ),
                                "metadata": metadatas_list[i],
                            }
                        )
                    except Exception:
                        continue

            print(f"Dense retrieval returned {len(dense_results)} results")
        except Exception as e:
            print(f"Error in dense retrieval: {e}")
            dense_results = []

        # Step 2: BM25 retrieval
        bm25_results = self._bm25_search(enhanced_query, top_k=initial_limit)
        print(f"BM25 retrieval returned {len(bm25_results)} results")

        # Step 3: Combine using Reciprocal Rank Fusion
        if dense_results or bm25_results:
            fused_results = self._reciprocal_rank_fusion(dense_results, bm25_results)
            print(f"RRF combined results: {len(fused_results)} total")
            for i, result in enumerate(fused_results[:3]):
                print(
                    f"  {i+1}. RRF: {result['rrf_score']:.3f} (dense: {result['dense_rank']}, BM25: {result['bm25_rank']})"
                )
        else:
            print("No results from either dense or BM25 retrieval")
            return []

        # Step 4: Cross-encoder reranking (optional, high-quality refinement)
        reranked_results = self._cross_encoder_rerank(
            query, fused_results, top_k=top_k * 3
        )

        # Use reranked results for final processing with DHO prioritization
        enhanced_results = reranked_results[
            : top_k * 2
        ]  # Take top results after reranking

        # Apply structural document type prioritization
        dho_results = []
        gaso_results = []
        other_results = []

        for result in enhanced_results:
            doc_name = result.get("metadata", {}).get("document_name", "")
            if "DHO" in doc_name:
                dho_results.append(result)
            elif "GASO" in doc_name or "Gp Air Staff" in doc_name:
                gaso_results.append(result)
            else:
                other_results.append(result)

        # Sort each category by score, then combine in priority order
        dho_results.sort(key=lambda x: x.get("cross_encoder_score", x["rrf_score"]), reverse=True)
        gaso_results.sort(key=lambda x: x.get("cross_encoder_score", x["rrf_score"]), reverse=True)
        other_results.sort(key=lambda x: x.get("cross_encoder_score", x["rrf_score"]), reverse=True)

        # Combine in priority order: DHO, GASO, Others
        enhanced_results = dho_results + gaso_results + other_results

        # Group chunks by document and reconstruct ProcessedDocuments
        document_chunks: dict[str, Any] = {}

        for result in enhanced_results:
            try:
                md = result["metadata"]
                document_id = str(md.get("document_id", "unknown"))
                document_name = str(md.get("document_name", "unknown"))
                file_type = str(md.get("file_type", "application/octet-stream"))
                directory_path = str(md.get("directory_path", "."))
                chunk_content = str(result["content"])
            except Exception:
                continue

            if document_id not in document_chunks:
                document_chunks[document_id] = {
                    "name": document_name,
                    "file_type": file_type,
                    "directory_path": directory_path,
                    "chunks": [],
                    "full_content": "",
                    "max_score": result.get("cross_encoder_score", result["rrf_score"]),
                    "score_type": (
                        "cross_encoder" if "cross_encoder_score" in result else "rrf"
                    ),
                }
            else:
                # Update max score for document
                current_score = result.get("cross_encoder_score", result["rrf_score"])
                document_chunks[document_id]["max_score"] = max(
                    document_chunks[document_id]["max_score"], current_score
                )

            document_chunks[document_id]["chunks"].append(chunk_content)
            document_chunks[document_id]["full_content"] += chunk_content + " "

        # Reconstruct ProcessedDocument objects and sort by best available score
        processed_docs = []
        for doc_id, doc_data in document_chunks.items():
            processed_doc = ProcessedDocument(
                id=doc_id,
                original_document=Document(
                    name=doc_data["name"],
                    file_path=f"{doc_data['directory_path']}/{doc_data['name']}",
                    file_type=doc_data["file_type"],
                    directory_path=doc_data["directory_path"],
                ),
                content=doc_data["full_content"].strip(),
                chunks=doc_data["chunks"],
                processed_at=datetime.now(UTC),
            )
            # Store best score for sorting
            processed_doc._hybrid_score = doc_data["max_score"]  # type: ignore
            processed_doc._score_type = doc_data["score_type"]  # type: ignore
            processed_docs.append(processed_doc)

        # Sort by best available score and return top_k
        processed_docs.sort(key=lambda x: getattr(x, "_hybrid_score", 0), reverse=True)
        return processed_docs[:top_k]

    def _model_requires_prefixes(self, model_name: str) -> bool:
        """Check if the embedding model requires query/passage prefixes.

        Args:
            model_name: Name of the embedding model

        Returns:
            True if model requires prefixes (e.g., E5 models)
        """
        return "e5-" in model_name.lower()

    def _format_for_embedding(self, text: str, is_query: bool = False) -> str:
        """Format text for embedding based on model requirements.

        Args:
            text: Text to format
            is_query: True if this is a search query, False for document passage

        Returns:
            Formatted text with appropriate prefixes if needed
        """
        if not self._requires_prefixes:
            return text

        prefix = "query: " if is_query else "passage: "
        return prefix + text

    async def index_documents(self, processed_docs: list[ProcessedDocument]) -> None:
        """Index processed documents for fast retrieval.

        Args:
            processed_docs: List of processed documents to index
        """
        for doc in processed_docs:
            # Store document chunks with embeddings in ChromaDB
            if not doc.embeddings or len(doc.embeddings) != len(doc.chunks):
                print(
                    f"Warning: Document {doc.original_document.name} has invalid embeddings"
                )
                continue

            for i, (chunk, embedding) in enumerate(
                zip(doc.chunks, doc.embeddings, strict=False)
            ):
                chunk_id = f"{doc.id}_chunk_{i}"

                # Extract enhanced metadata from chunk content
                chunk_page_numbers = self._extract_chunk_page_numbers(chunk)
                chunk_section_title = self._extract_chunk_section_title(chunk)

                # New enhanced metadata
                chunk_type = self._detect_chunk_type(chunk)
                annex_info = self._detect_annex(chunk)
                table_title = self._extract_table_title(chunk)

                # Extract page start/end from page numbers
                page_parts = chunk_page_numbers.split("-")
                page_start = page_parts[0]
                page_end = page_parts[-1] if len(page_parts) > 1 else page_parts[0]

                # Get section path with document context
                all_lines = doc.content.split("\n")
                # Estimate chunk start line (approximate)
                chunk_start_line = sum(
                    len(prev_chunk.split("\n")) for prev_chunk in doc.chunks[:i]
                )
                section_path = self._extract_section_path(
                    chunk, all_lines, chunk_start_line
                )

                # Extract key terms for better searchability
                chunk_key_terms = self.retrieval_enhancer.extract_key_terms(chunk)

                self.collection.add(
                    embeddings=[embedding],
                    documents=[chunk],
                    metadatas=[
                        {
                            "document_id": doc.id,
                            "document_name": doc.original_document.name,
                            "chunk_index": i,
                            "file_type": doc.original_document.file_type,
                            "directory_path": doc.original_document.directory_path,
                            "page_numbers": chunk_page_numbers,
                            "page_start": page_start,
                            "page_end": page_end,
                            "section_title": chunk_section_title,
                            "section_path": section_path,
                            "chunk_type": chunk_type,
                            "annex": annex_info,
                            "table_title": table_title,
                            "key_terms": ";".join(
                                chunk_key_terms[:10]
                            ),  # Limit to 10 terms
                        }
                    ],
                    ids=[chunk_id],
                )

                # Add to BM25 corpus for hybrid retrieval
                self._bm25_corpus.append(chunk)
                self._bm25_ids.append(chunk_id)

            # Add document-level summary embedding if available
            if hasattr(doc, "_doc_summary_embedding") and hasattr(doc, "_doc_summary"):
                summary_id = f"{doc.id}_summary"
                summary_embedding = doc._doc_summary_embedding
                summary_text = doc._doc_summary

                self.collection.add(
                    embeddings=[summary_embedding],
                    documents=[summary_text],
                    metadatas=[
                        {
                            "document_id": doc.id,
                            "document_name": doc.original_document.name,
                            "chunk_index": -1,  # Special marker for document summary
                            "file_type": doc.original_document.file_type,
                            "directory_path": doc.original_document.directory_path,
                            "page_numbers": "all",
                            "page_start": "1",
                            "page_end": "all",
                            "section_title": "Document Summary",
                            "section_path": "Document Overview",
                            "chunk_type": "summary",
                            "annex": "",
                            "table_title": "",
                            "key_terms": "",
                        }
                    ],
                    ids=[summary_id],
                )

                # Add summary to BM25 corpus
                self._bm25_corpus.append(summary_text)
                self._bm25_ids.append(summary_id)

        total_chunks = sum(len(doc.chunks) for doc in processed_docs)
        summary_count = sum(
            1 for doc in processed_docs if hasattr(doc, "_doc_summary_embedding")
        )
        print(
            f"Indexed {len(processed_docs)} documents with {total_chunks} content chunks and {summary_count} document summaries"
        )

        # Build BM25 index for hybrid retrieval
        self._build_bm25_index()

        # Update manifest for each document
        for doc in processed_docs:
            chunk_count = len(doc.chunks) if doc.chunks else 0
            self.manifest.update_document_entry(
                doc.original_document, chunk_count, self.embedding_model_name
            )

    async def process_changed_documents(
        self, documents: list[Document]
    ) -> list[ProcessedDocument]:
        """Process only documents that have changed since last processing.

        Args:
            documents: List of documents to check for changes

        Returns:
            List of processed documents (only changed ones)
        """
        # Get only changed documents
        changed_docs = self.manifest.get_changed_documents(documents)

        if not changed_docs:
            print("No document changes detected - skipping processing")
            return []

        print(
            f"Processing {len(changed_docs)} changed documents out of {len(documents)} total"
        )
        for doc in changed_docs:
            print(f"  - {doc.name} (changed)")

        # Process only changed documents
        return await self.process_documents(changed_docs)

    def should_full_reindex(self) -> bool:
        """Check if full reindex is needed due to embedding model change.

        Returns:
            True if full reindex is needed
        """
        return self.manifest.should_full_reindex(self.embedding_model_name)

    def mark_full_reindex_complete(self) -> None:
        """Mark that a full reindex has been completed."""
        self.manifest.mark_full_reindex(self.embedding_model_name)

    def get_manifest_stats(self) -> dict[str, Any]:
        """Get manifest statistics for health monitoring.

        Returns:
            Dictionary with manifest statistics
        """
        return self.manifest.get_stats()

    def clear_manifest(self) -> None:
        """Clear manifest entries (use when doing full reindex)."""
        self.manifest.clear()

    def _extract_chunk_page_numbers(self, chunk: str) -> str:
        """Extract page numbers from a specific chunk.

        Args:
            chunk: Text chunk to analyze

        Returns:
            Page number range or single page
        """
        import re

        page_numbers = set()

        # Look for page markers in the chunk
        lines = chunk.split("\n")
        for line in lines:
            if "--- PAGE" in line:
                match = re.search(r"PAGE\s+(\d+)", line)
                if match:
                    page_numbers.add(match.group(1))

        if page_numbers:
            sorted_pages = sorted(page_numbers, key=int)
            if len(sorted_pages) == 1:
                return sorted_pages[0]
            else:
                return f"{sorted_pages[0]}-{sorted_pages[-1]}"

        return "1"  # Fallback

    def _extract_chunk_section_title(self, chunk: str) -> str:
        """Extract section title from a specific chunk using improved aviation patterns.

        Args:
            chunk: Text chunk to analyze

        Returns:
            Section title if found
        """
        lines = chunk.split("\n")

        # Look through the first few lines for a proper section title
        for line in lines[:5]:  # Reduced from 10 to focus on actual headers
            line_stripped = line.strip()
            if not line_stripped:
                continue

            # Check for enhanced section headers (=== ... ===)
            if line_stripped.startswith("===") and line_stripped.endswith("==="):
                section_title = line_stripped.replace("===", "").strip()
                if section_title:
                    return section_title

            # Check for markdown headers
            if line_stripped.startswith("#"):
                return line_stripped.lstrip("#").strip()

            # Check for aviation document references
            import re

            if re.match(
                r"^(DHO|ANNEX|APPENDIX|SECTION|CHAPTER|PART)\s+[A-Z0-9]",
                line_stripped.upper(),
            ):
                return line_stripped

            # Check for structured aviation section titles (must be at start of chunk)
            if line == lines[0]:  # Only check first line to avoid random matches
                if (
                    line_stripped.isupper()
                    and 10 <= len(line_stripped) <= 80
                    and any(
                        word in line_stripped
                        for word in [
                            "WEATHER",
                            "OPERATIONS",
                            "LIMITS",
                            "PROCEDURES",
                            "SAFETY",
                            "EMERGENCY",
                        ]
                    )
                ):
                    return line_stripped

        # Try to extract from page markers or document structure
        for line in lines[:3]:
            if line.strip().startswith("--- PAGE"):
                return f"Page Content ({line.strip()})"

        # If we find specific aviation content, use more descriptive titles
        chunk_lower = chunk.lower()
        if "weather" in chunk_lower and (
            "limit" in chunk_lower or "wind" in chunk_lower
        ):
            return "Weather Limitations"
        elif "emergency" in chunk_lower and "procedure" in chunk_lower:
            return "Emergency Procedures"
        elif "operational" in chunk_lower and (
            "limit" in chunk_lower or "procedure" in chunk_lower
        ):
            return "Operational Procedures"

        return "General Content"

    def _detect_chunk_type(self, chunk: str) -> str:
        """Detect the type of content in a chunk.

        Args:
            chunk: Text chunk to analyze

        Returns:
            Chunk type classification
        """
        chunk_lower = chunk.lower()

        # Check for tables
        if (
            "table" in chunk_lower and any(marker in chunk for marker in ["|", "---"])
        ) or chunk.count("|") > 3:
            return "table"

        # Check for lists
        if (
            chunk.count("•") > 2
            or chunk.count("-") > 3
            or any(chunk.strip().startswith(f"{i}.") for i in range(1, 10))
        ):
            return "list"

        # Check for procedures
        if any(
            word in chunk_lower for word in ["procedure", "step", "action", "follow"]
        ) and ("1." in chunk or "a)" in chunk):
            return "procedure"

        # Check for headers/titles
        lines = chunk.split("\n")
        if len(lines) <= 3 and any(
            line.isupper() or line.startswith("===") for line in lines
        ):
            return "header"

        # Check for weather/limits tables
        if any(
            term in chunk_lower
            for term in ["wind", "crosswind", "knots", "kt", "weather limit"]
        ):
            return "weather_limits"

        return "text"

    def _detect_annex(self, chunk: str) -> str:
        """Detect if chunk is part of an annex.

        Args:
            chunk: Text chunk to analyze

        Returns:
            Annex identifier or empty string
        """
        import re

        # Look for annex references
        annex_match = re.search(r"annex\s+([a-z])", chunk.lower())
        if annex_match:
            return f"Annex {annex_match.group(1).upper()}"

        # Look for appendix references
        appendix_match = re.search(r"appendix\s+([a-z0-9]+)", chunk.lower())
        if appendix_match:
            return f"Appendix {appendix_match.group(1).upper()}"

        return ""

    def _extract_table_title(self, chunk: str) -> str:
        """Extract table title if chunk contains a table.

        Args:
            chunk: Text chunk to analyze

        Returns:
            Table title or empty string
        """
        lines = chunk.split("\n")

        # Look for table markers and nearby titles
        for line in lines:
            if "table" in line.lower() and (":" in line or len(line.split()) <= 8):
                return line.strip()
            elif "--- table" in line.lower():
                # Extract table number/title
                import re

                match = re.search(r"table\s+(\d+|[a-z])", line.lower())
                if match:
                    return f"Table {match.group(1).upper()}"

        # Look for weather limitations table specifically
        if "weather" in chunk.lower() and "limit" in chunk.lower():
            return "Weather Limitations Table"

        return ""

    def _extract_section_path(
        self, chunk: str, all_lines: list[str], start_line: int
    ) -> str:
        """Extract hierarchical section path for a chunk.

        Args:
            chunk: Text chunk to analyze
            all_lines: All document lines for context
            start_line: Starting line number of chunk in document

        Returns:
            Hierarchical section path
        """
        # Use existing header path extraction but make it more structured
        header_path = self._extract_header_path(all_lines, start_line)

        # Also check chunk content for additional context
        chunk_lines = chunk.split("\n")
        for line in chunk_lines[:3]:
            line_stripped = line.strip()
            if line_stripped.startswith("===") and line_stripped.endswith("==="):
                section_title = line_stripped.replace("===", "").strip()
                if section_title and section_title not in header_path:
                    header_path = (
                        f"{header_path} > {section_title}"
                        if header_path
                        else section_title
                    )
                break

        return header_path

    def _generate_document_summary(self, content: str, document_name: str) -> str:
        """Generate a concise summary of the document for whole-document queries.

        Args:
            content: Full document content
            document_name: Name of the document

        Returns:
            Document summary paragraph
        """
        # Extract key sections and metadata for summary
        lines = content.split("\n")
        key_sections = []
        key_terms = set()

        # Extract section headings and important content
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue

            # Section headers
            if (
                line_stripped.startswith("===") and line_stripped.endswith("===")
            ) or line_stripped.startswith("#"):
                section_title = line_stripped.replace("===", "").lstrip("#").strip()
                if section_title and len(section_title) < 100:
                    key_sections.append(section_title)

            # Aviation-specific terms
            line_lower = line_stripped.lower()
            for term in [
                "weather",
                "wind",
                "operational",
                "safety",
                "emergency",
                "procedures",
                "limits",
            ]:
                if term in line_lower:
                    key_terms.add(term)

        # Create summary
        summary_parts = [f"Document: {document_name}"]

        if key_sections:
            # Limit to most important sections
            important_sections = key_sections[:5]
            summary_parts.append(f"Key sections: {', '.join(important_sections)}")

        if key_terms:
            summary_parts.append(f"Topics covered: {', '.join(sorted(key_terms))}")

        # Add content type detection
        content_lower = content.lower()
        content_types = []
        if "table" in content_lower and any(
            marker in content for marker in ["|", "---"]
        ):
            content_types.append("tables")
        if content.count("•") > 5 or content.count("-") > 10:
            content_types.append("lists")
        if any(term in content_lower for term in ["procedure", "step", "action"]):
            content_types.append("procedures")
        if any(
            term in content_lower for term in ["weather limit", "wind", "crosswind"]
        ):
            content_types.append("weather limitations")

        if content_types:
            summary_parts.append(f"Contains: {', '.join(content_types)}")

        return ". ".join(summary_parts) + "."

    def _build_bm25_index(self) -> None:
        """Build BM25 index from stored corpus."""
        try:
            # Try to import BM25, install if not available
            try:
                from rank_bm25 import BM25Okapi
            except ImportError:
                print("BM25 library not found. Installing rank-bm25...")
                import subprocess
                import sys

                subprocess.check_call(
                    [sys.executable, "-m", "pip", "install", "rank-bm25"]
                )
                from rank_bm25 import BM25Okapi

            if self._bm25_corpus:
                # Tokenize corpus for BM25
                tokenized_corpus = [doc.lower().split() for doc in self._bm25_corpus]
                self._bm25_index = BM25Okapi(tokenized_corpus)
                print(f"Built BM25 index with {len(self._bm25_corpus)} documents")
            else:
                print("No corpus available for BM25 indexing")
        except Exception as e:
            print(f"Error building BM25 index: {e}")
            self._bm25_index = None

    def _bm25_search(self, query: str, top_k: int = 50) -> list[dict[str, Any]]:
        """Perform BM25 search on the corpus.

        Args:
            query: Search query
            top_k: Number of top results to return

        Returns:
            List of BM25 search results with scores
        """
        if not self._bm25_index or not self._bm25_corpus:
            return []

        try:
            query_tokens = query.lower().split()
            scores = self._bm25_index.get_scores(query_tokens)

            # Get top-k results with scores
            scored_results = [(score, i) for i, score in enumerate(scores)]
            scored_results.sort(reverse=True, key=lambda x: x[0])

            results = []
            for score, idx in scored_results[:top_k]:
                if idx < len(self._bm25_ids) and idx < len(self._bm25_corpus):
                    results.append(
                        {
                            "chunk_id": self._bm25_ids[idx],
                            "content": self._bm25_corpus[idx],
                            "bm25_score": float(score),
                            "rank": len(results) + 1,
                        }
                    )

            return results
        except Exception as e:
            print(f"Error in BM25 search: {e}")
            return []

    def _reciprocal_rank_fusion(
        self,
        dense_results: list[dict[str, Any]],
        bm25_results: list[dict[str, Any]],
        k: int = 60,
    ) -> list[dict[str, Any]]:
        """Combine dense and BM25 results using Reciprocal Rank Fusion.

        Args:
            dense_results: Results from dense retrieval
            bm25_results: Results from BM25 retrieval
            k: RRF parameter (default 60)

        Returns:
            Fused results sorted by RRF score
        """
        # Create mapping of chunk_id to scores
        rrf_scores: dict[str, float] = {}
        chunk_data: dict[str, dict[str, Any]] = {}

        # Add dense results
        for rank, result in enumerate(dense_results, 1):
            chunk_id = result["chunk_id"]
            rrf_scores[chunk_id] = rrf_scores.get(chunk_id, 0) + 1.0 / (k + rank)
            chunk_data[chunk_id] = result
            chunk_data[chunk_id]["dense_rank"] = rank

        # Add BM25 results
        for rank, result in enumerate(bm25_results, 1):
            chunk_id = result["chunk_id"]
            rrf_scores[chunk_id] = rrf_scores.get(chunk_id, 0) + 1.0 / (k + rank)
            if chunk_id not in chunk_data:
                chunk_data[chunk_id] = result
            chunk_data[chunk_id]["bm25_rank"] = rank
            chunk_data[chunk_id]["bm25_score"] = result.get("bm25_score", 0.0)

        # Sort by RRF score
        sorted_results = []
        for chunk_id, rrf_score in sorted(
            rrf_scores.items(), key=lambda x: x[1], reverse=True
        ):
            result = chunk_data[chunk_id].copy()
            result["rrf_score"] = rrf_score
            result["dense_rank"] = result.get("dense_rank", 999)
            result["bm25_rank"] = result.get("bm25_rank", 999)
            sorted_results.append(result)

        return sorted_results

    def _cross_encoder_rerank(
        self, query: str, results: list[dict[str, Any]], top_k: int = 12
    ) -> list[dict[str, Any]]:
        """Apply cross-encoder reranking to refine results.

        Args:
            query: Original search query
            results: Results to rerank
            top_k: Number of top results to keep after reranking

        Returns:
            Reranked results with cross-encoder scores
        """
        try:
            # Try to import cross-encoder, install if not available
            try:
                from sentence_transformers import CrossEncoder
            except ImportError:
                print("CrossEncoder not available, skipping reranking")
                return results[:top_k]

            # Try to load a cross-encoder model
            try:
                cross_encoder = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2")
                print("Loaded cross-encoder: ms-marco-MiniLM-L-6-v2")
            except Exception:
                try:
                    cross_encoder = CrossEncoder("BAAI/bge-reranker-base")
                    print("Loaded cross-encoder: bge-reranker-base")
                except Exception:
                    print("Could not load cross-encoder, skipping reranking")
                    return results[:top_k]

            # Prepare query-document pairs for reranking
            pairs = []
            for result in results[:50]:  # Limit to top 50 for reranking
                content = result["content"]
                # Truncate content if too long
                if len(content) > 512:
                    content = content[:512] + "..."
                pairs.append([query, content])

            if not pairs:
                return results[:top_k]

            # Get cross-encoder scores
            scores = cross_encoder.predict(pairs)

            # Add scores to results and sort
            reranked_results = []
            for i, result in enumerate(results[: len(scores)]):
                result_copy = result.copy()
                result_copy["cross_encoder_score"] = float(scores[i])
                reranked_results.append(result_copy)

            # Sort by cross-encoder score
            reranked_results.sort(key=lambda x: x["cross_encoder_score"], reverse=True)

            print(f"Cross-encoder reranked {len(reranked_results)} results")
            for i, result in enumerate(reranked_results[:3]):
                print(
                    f"  {i+1}. Cross-encoder: {result['cross_encoder_score']:.3f} (RRF: {result.get('rrf_score', 0):.3f})"
                )

            return reranked_results[:top_k]

        except Exception as e:
            print(f"Error in cross-encoder reranking: {e}")
            return results[:top_k]

    def _validate_and_filter_chunks(
        self, chunks: list[str], document_name: str
    ) -> list[str]:
        """Validate and filter chunks for safety and quality.

        Args:
            chunks: List of text chunks to validate
            document_name: Name of source document for logging

        Returns:
            List of validated chunks
        """
        validated_chunks = []
        empty_count = 0
        oversized_count = 0
        duplicate_count = 0
        max_chunk_tokens = 1500  # Allow larger chunks for complete aviation sections

        for i, chunk in enumerate(chunks):
            chunk_stripped = chunk.strip()

            # Skip empty chunks
            if not chunk_stripped:
                empty_count += 1
                continue

            # Skip chunks that are too short to be meaningful
            if len(chunk_stripped) < 10:
                empty_count += 1
                continue

            # Skip duplicate chunks
            if self._is_duplicate_chunk(chunk_stripped):
                duplicate_count += 1
                continue

            # Check chunk size using accurate token counting
            chunk_tokens = self._token_count(chunk_stripped)
            if chunk_tokens > max_chunk_tokens:
                oversized_count += 1
                print(
                    f"Warning: Chunk {i} in {document_name} exceeds size limit ({chunk_tokens} tokens), truncating"
                )
                # Truncate at sentence boundary if possible
                sentences = chunk_stripped.split(". ")
                truncated = ""
                for sentence in sentences:
                    test_truncated = truncated + sentence + ". "
                    if self._token_count(test_truncated) < max_chunk_tokens:
                        truncated = test_truncated
                    else:
                        break
                # If sentence-based truncation failed, use token-based truncation
                if not truncated.strip():
                    tokens = self.tokenizer.encode(chunk_stripped)
                    truncated_tokens = tokens[:max_chunk_tokens]
                    chunk_stripped = self.tokenizer.decode(truncated_tokens)
                else:
                    chunk_stripped = truncated.strip()

            # Additional quality checks
            if self._is_chunk_meaningful(chunk_stripped):
                validated_chunks.append(chunk_stripped)

        # Log validation results
        if empty_count > 0:
            print(f"Filtered out {empty_count} empty/short chunks from {document_name}")
        if duplicate_count > 0:
            print(
                f"Filtered out {duplicate_count} duplicate chunks from {document_name}"
            )
        if oversized_count > 0:
            print(f"Truncated {oversized_count} oversized chunks from {document_name}")

        print(
            f"Validated {len(validated_chunks)} chunks from {len(chunks)} original chunks for {document_name}"
        )
        return validated_chunks

    def _is_chunk_meaningful(self, chunk: str) -> bool:
        """Check if chunk contains meaningful content.

        Args:
            chunk: Text chunk to check

        Returns:
            True if chunk is meaningful
        """
        # Check for minimum content requirements
        words = chunk.split()
        if len(words) < 3:
            return False

        # Check that chunk isn't just repetitive content
        unique_words = {word.lower() for word in words}
        if len(unique_words) < len(words) * 0.3:  # Less than 30% unique words
            return False

        # Check for at least some alphabetic content
        alpha_chars = sum(1 for char in chunk if char.isalpha())
        if alpha_chars < len(chunk) * 0.5:  # Less than 50% alphabetic
            return False

        return True


if __name__ == "__main__":
    import argparse
    import asyncio
    from pathlib import Path

    from vgs_chatbot.models.document import Document

    parser = argparse.ArgumentParser(
        description="Reprocess and reindex all documents for embeddings."
    )
    parser.add_argument(
        "--documents-dir",
        type=str,
        default="data/documents",
        help="Directory containing documents to process (default: data/documents)",
    )
    parser.add_argument(
        "--persist-dir",
        type=str,
        default="data/vectors/chroma",
        help="Directory for persistent ChromaDB storage (default: data/vectors/chroma)",
    )
    parser.add_argument(
        "--embedding-model",
        type=str,
        default="BAAI/bge-small-en-v1.5",
        help="SentenceTransformer model to use (default: BAAI/bge-small-en-v1.5)",
    )
    args = parser.parse_args()

    documents_dir = Path(args.documents_dir)
    persist_dir = args.persist_dir
    embedding_model = args.embedding_model

    # Remove existing ChromaDB index if present
    chroma_path = Path(persist_dir)
    if chroma_path.exists():
        import shutil

        print(f"Removing existing ChromaDB index at {chroma_path}...")
        shutil.rmtree(chroma_path)

    # Gather all documents
    doc_files = [f for f in documents_dir.glob("*") if f.is_file()]
    if not doc_files:
        print(f"No documents found in {documents_dir}. Exiting.")
        exit(1)

    # Guess file types
    EXT_TO_MIME = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xls": "application/vnd.ms-excel",
        ".txt": "text/plain",
        ".md": "text/markdown",
    }
    documents = []
    for f in doc_files:
        ext = f.suffix.lower()
        file_type = EXT_TO_MIME.get(ext, "application/octet-stream")
        documents.append(
            Document(
                name=f.name,
                file_path=str(f),
                file_type=file_type,
                directory_path=str(f.parent),
            )
        )

    print(
        f"Found {len(documents)} documents. Processing and indexing with model '{embedding_model}'..."
    )

    processor = RAGDocumentProcessor(
        embedding_model=embedding_model,
        persist_directory=persist_dir,
    )

    async def process_and_index() -> None:
        processed = await processor.process_documents(documents)
        await processor.index_documents(processed)

    asyncio.run(process_and_index())
    print("✅ Reprocessing and reindexing complete.")
